"""
解析模块单元测试

测试覆盖：
1. PDF 纯文本提取（pdfplumber）
2. DOCX 结构提取（标题、段落、表格、图片）
3. 解析器包统一入口
4. 不支持格式的异常处理

运行方式: pytest backend/tests/test_parsing.py -v
"""

import logging

logger = logging.getLogger(__name__)

import io
import pytest

from knowlebase.parsers.base import ParsedSection, ParsedTable
from knowlebase.parsers.pdf_parser import PDFParser
from knowlebase.parsers.docx_parser import DOCXParser
from knowlebase.parsers import parse_document


# ============ 测试 1: PDF 纯文本提取 ============

@pytest.mark.asyncio
async def test_pdf_text_extraction():
    """PDFParser 能提取纯文本"""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_buf = io.BytesIO()
    c = canvas.Canvas(pdf_buf, pagesize=letter)
    c.drawString(100, 700, "This is a test document")
    c.drawString(100, 650, "Second line with numbers: 12345")
    c.drawString(100, 600, "Third line: mixed content here")
    c.save()
    pdf_bytes = pdf_buf.getvalue()

    parser = PDFParser()
    result = await parser.parse(pdf_bytes)

    assert result.page_count == 1, f"期望 page_count=1, 实际={result.page_count}"
    assert len(result.sections) == 1, f"期望 1 个 section, 实际={len(result.sections)}"

    all_text = ""
    for item in result.sections[0].content:
        all_text += item.text
        assert item.page_number == 1, f"期望 page_number=1, 实际={item.page_number}"
        assert item.type == "text", f"期望 type=text, 实际={item.type}"

    assert "test document" in all_text, "文本中应包含'test document'"
    assert "Second line" in all_text, "文本中应包含'Second line'"


@pytest.mark.asyncio
async def test_pdf_blank_page():
    """PDFParser 处理空白页面"""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.save()

    parser = PDFParser()
    result = await parser.parse(buf.getvalue())

    # reportlab 创建的空 PDF 有 0 个页面
    assert result.page_count == 0
    assert len(result.sections) == 1
    assert len(result.sections[0].content) == 0


# ============ 测试 2: DOCX 结构提取 ============

@pytest.mark.asyncio
async def test_docx_paragraph_extraction():
    """DOCXParser 能提取段落文本"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("第一段测试文本")
    doc.add_paragraph("第二段：包含中文和 English")

    buf = io.BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()

    parser = DOCXParser()
    result = await parser.parse(doc_bytes)

    assert len(result.sections) == 1, f"期望 1 个 section, 实际={len(result.sections)}"
    content = result.sections[0].content
    text_items = [item for item in content if item.type == "text"]
    assert len(text_items) >= 1, "应至少提取一个文本段落"
    assert text_items[0].type == "text"


@pytest.mark.asyncio
async def test_docx_heading_detection():
    """DOCXParser 能识别标题"""
    from docx import Document

    doc = Document()
    doc.add_heading("第一章 概述", level=1)
    doc.add_paragraph("这是第一章的内容")
    doc.add_heading("1.1 小节", level=2)
    doc.add_paragraph("这是小节内容")

    buf = io.BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()

    parser = DOCXParser()
    result = await parser.parse(doc_bytes)

    content = result.sections[0].content
    sections = [item for item in content if item.type == "section"]
    assert len(sections) >= 1, "应至少识别一个标题 section"
    section_titles = [s.title for s in sections]
    assert "第一章 概述" in section_titles, f"应识别到'第一章 概述'，实际: {section_titles}"


@pytest.mark.asyncio
async def test_docx_table_extraction():
    """DOCXParser 能提取表格"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("表格的上方文本")

    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "姓名"
    table.rows[0].cells[1].text = "年龄"
    table.rows[1].cells[0].text = "张三"
    table.rows[1].cells[1].text = "25"
    table.rows[2].cells[0].text = "李四"
    table.rows[2].cells[1].text = "30"

    buf = io.BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()

    parser = DOCXParser()
    result = await parser.parse(doc_bytes)

    content = result.sections[0].content
    tables = [item for item in content if isinstance(item, ParsedTable)]
    assert len(tables) >= 1, "应至少提取一个表格"
    assert tables[0].headers == ["姓名", "年龄"], f"表头应为 ['姓名', '年龄'], 实际={tables[0].headers}"
    assert len(tables[0].rows) == 2, f"应有 2 行数据, 实际={len(tables[0].rows)}"


# ============ 测试 3: 统一入口 ============

@pytest.mark.asyncio
async def test_parse_document_router():
    """parse_document 根据文件名自动选择解析器"""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from docx import Document

    # PDF 解析器
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(100, 700, "test content")
    c.save()
    result = await parse_document(buf.getvalue(), filename="test.pdf")
    assert isinstance(result.sections, list)
    assert result.page_count == 1

    # DOCX 解析器
    doc = Document()
    doc.add_paragraph("test")
    buf2 = io.BytesIO()
    doc.save(buf2)
    result = await parse_document(buf2.getvalue(), filename="test.docx")
    assert len(result.sections) == 1


@pytest.mark.asyncio
async def test_parse_document_unsupported():
    """parse_document 不支持的格式抛出 ValueError"""
    with pytest.raises(ValueError, match="不支持"):
        await parse_document(b"fake content", filename="test.txt")


@pytest.mark.asyncio
async def test_parse_document_no_filename():
    """parse_document 不提供文件名和 mime_type 时抛出 ValueError"""
    with pytest.raises(ValueError):
        await parse_document(b"fake content")
