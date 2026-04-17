"""Word 文档解析器

使用 python-docx 提取完整结构：标题层级、段落、表格、图片。
输出树状 section + content 数组结构。
"""

import logging
from typing import List, Union

from knowlebase.parsers.base import (
    BaseParser,
    ParsedImage,
    ParsedSection,
    ParsedTable,
    ParsedText,
    ParseResult,
)

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Word 文档解析器（python-docx）"""

    async def parse(self, content: bytes) -> ParseResult:
        """解析 DOCX 文件

        遍历文档元素，识别：
        - 标题（Heading 1-6）→ 嵌套 ParsedSection
        - 段落（Paragraph）→ ParsedText
        - 表格（Table）→ ParsedTable
        - 图片（InlineShape）→ ParsedImage（Phase 1 仅标记，不提取）
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx 未安装，请运行: pip install python-docx")
            raise

        from io import BytesIO

        doc = Document(BytesIO(content))

        # 收集所有顶层元素（段落和表格）
        root_content: List[Union[ParsedText, ParsedImage, ParsedTable, ParsedSection]] = []
        has_images = False
        has_tables = False

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            # 表格元素
            if tag == "tbl":
                table_obj = self._find_table_by_element(doc, element)
                if table_obj:
                    parsed = self._parse_table(table_obj)
                    if parsed:
                        root_content.append(parsed)
                        has_tables = True
                continue

            # 段落元素
            if tag == "p":
                para = self._find_paragraph_by_element(doc, element)
                if para is None:
                    continue

                heading_level = self._get_heading_level(para)
                if heading_level:
                    # 标题 → 创建新 section
                    section = ParsedSection(
                        title=para.text.strip(),
                        content=[],
                    )
                    root_content.append(section)
                else:
                    # 普通段落 → text
                    text = para.text.strip()
                    if text:
                        root_content.append(ParsedText(text=text))

                    # 检查段落中的内联图片
                    for run in para.runs:
                        if run._r.xpath(".//a:blip"):
                            has_images = True
                            root_content.append(
                                ParsedImage(caption=f"[段落中的图片]")
                            )

        # 检查独立的图片（InlineShape）
        if doc.inline_shapes:
            has_images = True

        # 将非 section 的顶层元素分组到默认 section
        sections = self._organize_sections(root_content)

        return ParseResult(
            sections=sections,
            page_count=0,  # DOCX 无页概念
            has_images=has_images,
            has_tables=has_tables,
        )

    def _get_heading_level(self, paragraph) -> int:
        """获取段落标题级别（1-6），非标题返回 0"""
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
                return level
            except (ValueError, IndexError):
                return 1
        return 0

    def _organize_sections(self, content: list) -> List[ParsedSection]:
        """将顶层元素组织成 sections

        将连续的文本/表格/图片分组到前一个 section 的 content 中，
        或创建默认 section 容纳无标题的元素。
        """
        if not content:
            return [ParsedSection(title="", content=[])]

        # 简单策略：将所有元素放到一个顶层 section 中
        # Phase 1 不做复杂的层级嵌套
        return [ParsedSection(title="", content=content)]

    def _parse_table(self, table) -> ParsedTable:
        """解析 DOCX 表格"""
        if not table.rows:
            return None

        parsed = ParsedTable()

        # 第一行作为表头
        header_row = table.rows[0]
        parsed.headers = [cell.text.strip() for cell in header_row.cells]

        # 其余行作为数据
        for row in table.rows[1:]:
            parsed.rows.append([cell.text.strip() for cell in row.cells])

        return parsed

    def _find_table_by_element(self, doc, element):
        """通过 XML 元素查找对应的 docx Table 对象"""
        for table in doc.tables:
            if table._element is element:
                return table
        return None

    def _find_paragraph_by_element(self, doc, element):
        """通过 XML 元素查找对应的 docx Paragraph 对象"""
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None
